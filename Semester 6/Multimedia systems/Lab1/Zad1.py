import numpy as np
import matplotlib.pyplot as plt
import soundfile as sf
import scipy.fftpack
import scipy.fft
import sounddevice as sd
import math
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Zadanie 1
# data, fs = sf.read('sound1.wav', dtype='float32')
# left_chanel = data[:,0]
# right_chanel = data[:,1]
# mono = (data[:,0] + data[:,1])/2
# print(left_chanel.shape)
# print(right_chanel.shape)
# print(mono.shape)

# sf.write('sound_L.wav',left_chanel,fs)
# sf.write('sound_R.wav',right_chanel,fs)
# sf.write('sound_mix.wav',mono,fs)

# plt.subplot(2,1,1)
# plt.plot(data[:,0])
# plt.subplot(2,1,2)
# plt.plot(left_chanel)
# plt.show()



def plotAudio(signal, fs, time_margin=[0,0.02], fsize=2**8):
    plt.subplot(2,1,1)
    plt.plot(np.arange(0,signal.shape[0])/fs, signal)
    plt.xlim(time_margin)
    # widmo
    plt.subplot(2,1,2)
    yf = scipy.fftpack.fft(signal,fsize)
    plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
    plt.show()
# data, fs = sf.read('sin_440Hz.wav', dtype='float32')
# plotAudio(data,fs)
# data, fs = sf.read('sin_440Hz.wav', dtype='float32')
# plotAudio(data,fs, fsize=2**12)
# data, fs = sf.read('sin_440Hz.wav', dtype='float32')
# plotAudio(data,fs, fsize=2**16)


document = Document()
document.add_heading('lab1',0) # tworzenie nagłówków druga wartość to poziom nagłówka 


files=['sin_60Hz.wav','sin_440Hz.wav','sin_8000Hz.wav']
Margins=[2**8,2**12,2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file),2)
    for i,Margin in enumerate(Margins):
        document.add_heading('Time margin {}'.format(Margin),3) # nagłówek sekcji, mozę być poziom wyżej
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
    
        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        ############################################################
        time_margin = [0,0.02]
        signal, fs = sf.read(file, dtype='float32')
        fsize = Margins[i]
        plt.subplot(2,1,1)
        plt.plot(np.arange(0,signal.shape[0])/fs, signal)
        plt.xlim(time_margin)
        # widmo
        plt.subplot(2,1,2)
        yf = scipy.fftpack.fft(signal,fsize)
        plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
        plt.show()

        fig.suptitle('fsize {}'.format(Margin)) # Tytuł wykresu
        fig.tight_layout(pad=1.5) # poprawa czytelności 
        memfile = BytesIO() # tworzenie bufora
        fig.savefig(memfile) # z zapis do bufora 
        
    
        document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
        
        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph('Max wartość widma = {}'.format(np.argmax(yf))) 
        ############################################################

document.save('raport.docx') # zapis do pliku