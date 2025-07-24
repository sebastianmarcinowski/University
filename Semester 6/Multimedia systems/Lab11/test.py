import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
import matplotlib
from io import BytesIO
from docx import Document
from docx.shared import Inches
import pandas as pd
from skimage.metrics import structural_similarity as ssim
import random
import pandas as pd
import csv

matplotlib.use('TkAgg')


def water_mark(img,mask,alpha=0.25):
    assert (img.shape[0]==mask.shape[0]) and (img.shape[1]==mask.shape[1]), "Wrong size"
    if len(img.shape)<3:
        flag=True
        t_img=cv2.cvtColor(img,cv2.COLOR_GRAY2RGBA)
    else:
        flag=False
        t_img=cv2.cvtColor(img,cv2.COLOR_RGB2RGBA)
    if (mask.dtype==bool):
        t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
    elif (mask.dtype==np.uint8):
        if len(mask.shape)<3:
            t_mask=cv2.cvtColor((mask).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
        else:
            t_mask=cv2.cvtColor((mask).astype(np.uint8),cv2.COLOR_RGB2RGBA)
    else:
        if len(mask.shape)<3:
            t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
        else:
            t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_RGB2RGBA)
    t_out=cv2.addWeighted(t_img,1,t_mask,alpha,0)
    if flag:
        out=cv2.cvtColor(t_out,cv2.COLOR_RGBA2GRAY)
    else:
        out=cv2.cvtColor(t_out,cv2.COLOR_RGBA2RGB)
    return out



def put_data(img, data, binary_mask=np.uint8(1)):
    assert img.dtype == np.uint8, "img wrong data type"
    assert binary_mask.dtype == np.uint8, "binary_mask wrong data type"

    un_binary_mask = np.unpackbits(binary_mask)

    if data.dtype != bool:
        unpacked_data = np.unpackbits(data)
    else:
        unpacked_data = data

    # Calculate available space
    dataspace = img.shape[0] * img.shape[1] * np.sum(un_binary_mask)

    # Check if data fits
    if dataspace < unpacked_data.size:
        raise ValueError(f"Data too large! Need {unpacked_data.size} bits but only have {dataspace} bits available.")

    if dataspace == unpacked_data.size:
        prepared_data = unpacked_data.reshape(img.shape[0], img.shape[1], np.sum(un_binary_mask)).astype(np.uint8)
    else:
        padded_data = np.zeros(dataspace, dtype=np.uint8)
        padded_data[:unpacked_data.size] = unpacked_data
        prepared_data = padded_data.reshape(img.shape[0], img.shape[1], np.sum(un_binary_mask)).astype(np.uint8)

    mask = np.full((img.shape[0], img.shape[1]), binary_mask)
    img_copy = np.bitwise_and(img, np.invert(mask))

    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            temp = prepared_data[:, :, bv]
            temp = np.left_shift(temp, i)
            img_copy = np.bitwise_or(img_copy, temp)
            bv += 1

    return img_copy

def pop_data(img,binary_mask=np.uint8(1),out_shape=None):
    un_binary_mask=np.unpackbits(binary_mask)
    data=np.zeros((img.shape[0],img.shape[1],np.sum(un_binary_mask))).astype(np.uint8)
    bv=0
    for i,b in enumerate(un_binary_mask[::-1]):
        if b:
            mask=np.full((img.shape[0],img.shape[1]),2**i)
            temp=np.bitwise_and(img,mask)
            data[:,:,bv]=temp[:,:].astype(np.uint8)
            bv+=1
    if out_shape!=None:
        tmp=np.packbits(data.flatten())
        tmp=tmp[:np.prod(out_shape)]
        data=tmp.reshape(out_shape)
    return data


def MSE(original, compressed):
    mse = np.sum((original - compressed) ** 2) / original.size

    return mse


def PSNR(original, compressed):
    mse_val = MSE(original, compressed)
    if mse_val == 0:
        return 100.0
    if mse_val < 1e-10:
        return 100.0
    max_pixel_value = 255.0
    psnr = 20 * np.log10(max_pixel_value / np.sqrt(mse_val))
    return psnr

def SSIM(original, compressed):
    return ssim(original, compressed, win_size=3, channel_axis=-1)

def exercise2():
    document.add_heading('Zadanie 2. Kodowanie tekstu w obrazie', level=1)
    img = cv2.imread('color2.jpg')
    blue_channel = img[:, :, 0].copy()

    with open("tekst.txt", "r", encoding='utf-8') as file:
        text = file.read()

    binary_mask = np.uint8(1)
    text_data = np.frombuffer(text.encode('utf-8'), dtype=np.uint8)
    encoded_blue = put_data(blue_channel.copy(), text_data, binary_mask)

    restored_bits = pop_data(encoded_blue, binary_mask, out_shape=text_data.shape)
    restored_text = restored_bits.tobytes().decode('utf-8')

    document.add_paragraph(f"Oryginalny tekst:\n{text}")
    document.add_paragraph(f"Odzyskany tekst:\n{restored_text}")
    document.add_paragraph(f"Czy tekst odzyskany poprawnie?\n{text == restored_text}")

    psnr = PSNR(blue_channel, encoded_blue)
    ssim = SSIM(blue_channel, encoded_blue)

    document.add_paragraph(f"PSNR: {psnr:.2f} dB")
    document.add_paragraph(f"SSIM: {ssim:.4f}")

    document.add_heading('Wnioski', level=2)
    document.add_paragraph('Tekst został poprawnie zakodowany i odzyskany. Wartości PSNR i SSIM wskazują na wysoką jakość zakodowanego obrazu w porównaniu do oryginalnego obrazu.')

def exercise3():
    document.add_heading('Zadanie 3. Kodowanie obrazu w obrazie', level=1)
    img = cv2.imread('color1.jpg')
    if img is None:
        raise FileNotFoundError("Failed to load 'color1.jpg'. Check the file path or integrity.")
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

    hidden_img = cv2.imread('color2.jpg')
    if hidden_img is None:
        raise FileNotFoundError("Failed to load 'color2.png'. Check the file path or integrity.")
    hidden_img = cv2.cvtColor(hidden_img, cv2.COLOR_BGR2RGB)

    scale_factor = 0.5  # zmiejszanie rozmiaru obrazu ukrywanego, aby nie wyrzucało asercji
    while True:
        hidden_img = cv2.resize(hidden_img, (int(img.shape[1] * scale_factor), int(img.shape[0] * scale_factor)))
        hidden_b, hidden_g, hidden_r = cv2.split(hidden_img)

        carrier_b, carrier_g, carrier_r = cv2.split(img)
        dataspace_b = carrier_b.size * 3
        dataspace_rg = carrier_g.size * 2

        if hidden_b.size <= dataspace_b and hidden_g.size <= dataspace_rg and hidden_r.size <= dataspace_rg:
            break
        scale_factor *= 0.9

    binary_mask_b = np.uint8(7)
    binary_mask_rg = np.uint8(3)

    encoded_b = put_data(carrier_b, hidden_b, binary_mask_b)
    encoded_g = put_data(carrier_g, hidden_g, binary_mask_rg)
    encoded_r = put_data(carrier_r, hidden_r, binary_mask_rg)

    encoded_img = cv2.merge((encoded_b, encoded_g, encoded_r))

    restored_b = pop_data(encoded_b, binary_mask_b, out_shape=hidden_b.shape)
    restored_g = pop_data(encoded_g, binary_mask_rg, out_shape=hidden_g.shape)
    restored_r = pop_data(encoded_r, binary_mask_rg, out_shape=hidden_r.shape)
    restored_img = cv2.merge((restored_b, restored_g, restored_r))

    psnr_b = PSNR(hidden_b, restored_b)
    psnr_g = PSNR(hidden_g, restored_g)
    psnr_r = PSNR(hidden_r, restored_r)
    ssim_b = SSIM(hidden_b, restored_b)
    ssim_g = SSIM(hidden_g, restored_g)
    ssim_r = SSIM(hidden_r, restored_r)

    document.add_paragraph(f"PSNR (Blue): {psnr_b:.2f}, SSIM (Blue): {ssim_b:.4f}")
    document.add_paragraph(f"PSNR (Green): {psnr_g:.2f}, SSIM (Green): {ssim_g:.4f}")
    document.add_paragraph(f"PSNR (Red): {psnr_r:.2f}, SSIM (Red): {ssim_r:.4f}")


    fig, axs = plt.subplots(2, 2, figsize=(12, 8))
    axs[0, 0].imshow(img)
    axs[0, 0].set_title("Obraz nosiciel")
    axs[0, 1].imshow(encoded_img)
    axs[0, 1].set_title("Obraz zakodowany")
    axs[1, 0].imshow(hidden_img)
    axs[1, 0].set_title("Obraz ukryty")
    axs[1, 1].imshow(restored_img)
    axs[1, 1].set_title("Obraz odzyskany")
    for ax in axs.flat:
        ax.axis('off')
    plt.tight_layout()

    memfile = BytesIO()
    plt.savefig(memfile, format='png', bbox_inches='tight')
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(6))
    memfile.close()

    document.add_heading('Wnioski', level=2)
    document.add_paragraph('Na zakodowanym obrazie widać różnicę szczególnie w lewej części obrazu. Wartości PSNR i SSIM wskazują, że obraz ukryty i odzyskany są identyczne.')


def exercise4():

    img = cv2.imread('color2.jpg')
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

    bit_configs = [
        ("0 bits (no data)", 0, 0, 0, 0),
        ("1 bit (B only)", 0, 0, 1, 1),
        ("2 bits (RB)", 1, 0, 1, 2),
        ("3 bits (RGB)", 1, 1, 1, 3),
        ("4 bits (2R,2B)", 2, 0, 2, 4),
        ("5 bits (2R,1G,2B)", 2, 1, 2, 5),
        ("6 bits (2RGB)", 2, 2, 2, 6),
        ("7 bits (3R,2G,2B)", 3, 2, 2, 7),
        ("8 bits (3R,2G,3B)", 3, 2, 3, 8),
        ("9 bits (3RGB)", 3, 3, 3, 9),
        ("10 bits (4R,3G,3B)", 4, 3, 3, 10),
        ("11 bits (4R,4G,3B)", 4, 4, 3, 11),
        ("12 bits (4RGB)", 4, 4, 4, 12),
        ("13 bits (5R,4G,4B)", 5, 4, 4, 13),
        ("14 bits (5R,5G,4B)", 5, 5, 4, 14),
        ("15 bits (5RGB)", 5, 5, 5, 15),
        ("16 bits (6R,5G,5B)", 6, 5, 5, 16),
        ("17 bits (6R,6G,5B)", 6, 6, 5, 17),
        ("18 bits (6RGB)", 6, 6, 6, 18),
        ("19 bits (7R,6G,6B)", 7, 6, 6, 19),
        ("20 bits (7R,7G,6B)", 7, 7, 6, 20),
        ("21 bits (7RGB)", 7, 7, 7, 21),
        ("22 bits (8R,7G,7B)", 8, 7, 7, 22),
        ("23 bits (8R,8G,7B)", 8, 8, 7, 23),
        ("24 bits (8RGB)", 8, 8, 8, 24),
    ]

    results = []

    carrier_r, carrier_g, carrier_b = cv2.split(img)
    img_height, img_width = img.shape[0], img.shape[1]


    for desc, r_bits, g_bits, b_bits, total_bits in bit_configs:
        r_mask = np.uint8((2 ** r_bits) - 1) if r_bits > 0 else np.uint8(0)
        g_mask = np.uint8((2 ** g_bits) - 1) if g_bits > 0 else np.uint8(0)
        b_mask = np.uint8((2 ** b_bits) - 1) if b_bits > 0 else np.uint8(0)

        r_capacity_bits = img_height * img_width * r_bits
        g_capacity_bits = img_height * img_width * g_bits
        b_capacity_bits = img_height * img_width * b_bits

        random_data_r_bytes = np.ceil(r_capacity_bits / 8).astype(int)
        random_data_g_bytes = np.ceil(g_capacity_bits / 8).astype(int)
        random_data_b_bytes = np.ceil(b_capacity_bits / 8).astype(int)

        data_r = np.random.randint(0, 256, size=random_data_r_bytes, dtype=np.uint8) if r_bits > 0 else np.array([])
        data_g = np.random.randint(0, 256, size=random_data_g_bytes, dtype=np.uint8) if g_bits > 0 else np.array([])
        data_b = np.random.randint(0, 256, size=random_data_b_bytes, dtype=np.uint8) if b_bits > 0 else np.array([])

        encoded_r = put_data(carrier_r, data_r, r_mask) if r_bits > 0 else carrier_r
        encoded_g = put_data(carrier_g, data_g, g_mask) if g_bits > 0 else carrier_g
        encoded_b = put_data(carrier_b, data_b, b_mask) if b_bits > 0 else carrier_b

        encoded_img = cv2.merge((encoded_r, encoded_g, encoded_b))

        psnr_val = PSNR(img, encoded_img)
        ssim_val = SSIM(img, encoded_img)

        results.append({
            'Opis': desc,
            'Łączna liczba bitów': total_bits,
            'Bity R': r_bits,
            'Bity G': g_bits,
            'Bity B': b_bits,
            'PSNR': round(psnr_val, 2),
            'SSIM': round(ssim_val, 4),
            'Encoded_img': encoded_img.copy()
        })

    df_results = pd.DataFrame(results)
    df_results = df_results.drop('Encoded_img', axis=1)

    table = document.add_table(rows=1, cols=len(df_results.columns))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, column_name in enumerate(df_results.columns):
        header_cells[i].text = column_name

    for _, row in df_results.iterrows():
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = str(cell_value)


    for i in range(25):
        fix, axs = plt.subplots(1, 2, figsize=(12, 8))
        axs[0].imshow(img)
        axs[0].set_title("Obraz nosiciel")
        axs[1].imshow(results[i]['Encoded_img'])
        axs[1].set_title(f"{results[i]['Opis']}\nPSNR: {results[i]['PSNR']:.2f} dB, SSIM: {results[i]['SSIM']:.4f}")
        for ax in axs.flat:
            ax.axis('off')

        plt.tight_layout()

        memfile = BytesIO()
        plt.savefig(memfile, format='png', bbox_inches='tight')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()

    document.add_heading('Wnioski', level=2)
    document.add_paragraph(
        'Minimalny "budżet bitowy" moim zdaniem zaczyna się już od 10 bitów. Potwierdzają to wartości PSNR = 34.88 dB oraz SSIM = 0.8463, które wskazują na wyraźny spadek jakości względem niższych wartości, np. dla 6 bitów (PSNR = 44.15 dB, SSIM = 0.9723). '
        'W tej konfiguracji pojawiają się zauważalne różnice wizualne – przykładowo na koszulce można dostrzec zniekształcenia i różnice w kolorach.')

    document.add_paragraph(
        'Od 13 bitów (PSNR = 30.87 dB, SSIM = 0.6382) różnice są już mocno zauważalne – zarówno pod względem kontrastu, jak i deformacji detali. '
        'SSIM poniżej 0.7 oznacza, że struktura obrazu odbiega znacznie od oryginału, co przekłada się na pogorszenie percepcji.')

    document.add_paragraph(
        'Im dalej zwiększamy liczbę bitów, tym bardziej rośnie zniekształcenie. Przykładowo dla 18 bitów SSIM wynosi tylko 0.1995, a dla maksymalnych 24 bitów spada do 0.009. '
        'To oznacza niemal całkowitą utratę spójności z pierwotnym obrazem – ukryty obraz staje się zdominowany przez zakłócenia.')


def exercise5():
    document.add_heading('Zadanie 5. Wodny znak', level=1)
    img = cv2.imread('color2.jpg')
    if img is None:
        raise FileNotFoundError("Failed to load 'color2.jpeg'. Check the file path or integrity.")
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

    binary_mask = cv2.imread('binary1.jpg', cv2.IMREAD_GRAYSCALE)
    if binary_mask is None:
        raise FileNotFoundError("Failed to load 'binary1.png'. Check the file path or integrity.")
    binary_mask = cv2.resize(binary_mask, (img.shape[1], img.shape[0]))

    alpha_values = [0.10, 0.25, 0.50]

    for alpha in alpha_values:
        watermarked_img = water_mark(img, binary_mask, alpha=alpha)

        psnr_value = PSNR(img, watermarked_img)
        ssim_value = SSIM(img, watermarked_img)

        fix, axs = plt.subplots(1, 2, figsize=(12, 8))
        axs[0].imshow(img)
        axs[0].set_title("Obraz nosiciel")
        axs[1].imshow(watermarked_img)
        axs[1].set_title(f"Obraz z water markiem: alfa={alpha}")
        for ax in axs.flat:
            ax.axis('off')

        plt.tight_layout()

        memfile = BytesIO()
        plt.savefig(memfile, format='png', bbox_inches='tight')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close()  # Close the figure to prevent memory issues

        document.add_paragraph(f"PSNR: {psnr_value:.2f} dB, SSIM: {ssim_value:.4f}")

    document.add_heading('Wnioski', level=2)
    document.add_paragraph('Im wyższa wartość α, tym bardziej widoczny jest wodny znak.')
document = Document()
document.add_heading('Sprawozdanie z LAB11 - Łukasz Watral 53674', 0)
exercise2()
exercise3()
exercise4()
exercise5()
document.save('lab11_2.docx')