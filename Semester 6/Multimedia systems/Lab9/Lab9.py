import cv2
import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches

##############################################################################
######   Konfiguracja       ##################################################
##############################################################################

kat = 'wideo'  # katalog z plikami wideo
pliki = ["clip_3.mp4", "clip_4.mp4", "clip_5.mp4"]  # nazwa pliku
ile = 100  # ile klatek odtworzyć? <0 - całość
key_frame_counter = 4  # co która klatka ma być kluczowa i nie podlegać kompresji
key_frame_counters = [2, 4, 5, 6, 8, 9, 10]
plot_frames = np.array([47])  # automatycznie wyrysuj wykresy
auto_pause_frames = np.array([25])  # automatycznie za pauzuj dla klatki
subsampling = "4:2:2"  # parametry dla chroma subsampling
subsampling_rates = ['4:4:4', '4:2:2', '4:4:0', '4:2:0', '4:1:1', '4:1:0']  # dostępne opcje subsamplingu
dzielnik = 4  # dzielnik przy zapisie różnicy
dzielniki  = [1,2,4,8,16] # dostępne opcje dzielnika
wyswietlaj_kaltki = False # czy program ma wyświetlać klatki
ROI5 = [[50, 900, 50, 900]]  # wyświetlane fragmenty  - film nr 5
ROI = [[150, 800, 150, 800]]  # wyświetlane fragmenty  - film nr 3 i 4


##############################################################################
####     Kompresja i dekompresja    ##########################################
##############################################################################
class data:
    def init(self):
        self.Y = None
        self.Cb = None
        self.Cr = None

def count_same_symbols(obj):
    number_of_symbols = 1
    for i in range(len(obj)-1):
        if obj[i] == obj[i+1]:
            number_of_symbols += 1
        else:
            break
    return number_of_symbols

def count_diff_symbols(obj):
    number_of_symbols = 1
    for i in range(len(obj) - 1):
        if obj[i] == obj[i + 1]:
            number_of_symbols -= 1
            break
        else:
            number_of_symbols += 1
    return number_of_symbols

def ByteRun_encoder(obj):
    obj_temp = obj.copy()
    obj_temp = np.array(obj_temp)
    if obj_temp.shape != (1,):
        obj_temp = obj_temp.flatten()
    encoded = []
    i = 0
    while i < len(obj_temp):
        num_of_symbols = count_same_symbols(obj_temp[i:])
        if num_of_symbols > 1:
            encoded.append(-num_of_symbols + 1)
            encoded.append(obj_temp[i])
            i += num_of_symbols
        else:
            num_of_symbols = count_diff_symbols(obj_temp[i:])
            encoded.append(num_of_symbols - 1)
            encoded.extend(obj_temp[i:i + num_of_symbols])
            i += num_of_symbols
    return np.array(encoded, dtype=int)

def ByteRun_decoder(encoded):
    decoded = []
    i = 0
    while i < len(encoded):
        if encoded[i] < 0:  # Repeated symbols
            count = -encoded[i] + 1
            symbol = encoded[i + 1]
            decoded.extend([symbol] * count)
            i += 2
        else:  # Unique symbols
            count = encoded[i] + 1
            decoded.extend(encoded[i + 1:i + 1 + count])
            i += 1 + count
    return np.array(decoded, dtype=int).reshape(-1)

def no_encoder(data):
    return data

def no_decoder(data):
    return data

def Chroma_subsampling(L, subsampling):
    # uzupełnić
    L2 = L.copy()
    if subsampling=="4:2:2":
        L2 = L2[:, ::2]
    elif subsampling=="4:2:0":
        L2 = L2[::2, ::2]
    elif subsampling=="4:4:0":
        L2 = L2[::2,:]
    elif subsampling=="4:1:1":
        L2 = L2[:,::4]
    elif subsampling=="4:1:0":
        L2 = L2[::2, ::4]
    else: #defalut "4:4:4"
        pass
    return L2

def Chroma_resampling(L, subsampling):
    # uzupełnić
    L2 = L.copy()
    if subsampling=="4:2:2":
        L2 = np.repeat(L2, 2, axis=1)
    elif subsampling=="4:2:0":
        L2 = np.repeat(L2, 2, axis=0)
        L2 = np.repeat(L2, 2, axis=1)
    elif subsampling=="4:4:0":
        L2 = np.repeat(L2, 2, axis=0)
    elif subsampling=="4:1:1":
        L2 = np.repeat(L2, 4, axis=1)
    elif subsampling=="4:1:0":
        L2 = np.repeat(L2, 4, axis=1)
        L2 = np.repeat(L2, 2, axis=0)
    else: #defalut "4:4:4"
        pass
    return L2


def frame_image_to_class(frame, subsampling):
    Frame_class = data()
    Frame_class.Y = frame[:, :, 0].astype(int)
    Frame_class.Cb = Chroma_subsampling(frame[:, :, 2].astype(int), subsampling)
    Frame_class.Cr = Chroma_subsampling(frame[:, :, 1].astype(int), subsampling)
    return Frame_class


def frame_layers_to_image(Y, Cr, Cb, subsampling):
    Cb = Chroma_resampling(Cb, subsampling)
    Cr = Chroma_resampling(Cr, subsampling)
    return np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)


def compress_KeyFrame(Frame_class, encoder):
    KeyFrame = data()
    ## TO DO
    KeyFrame.Y = encoder(Frame_class.Y)
    KeyFrame.Cr = encoder(Frame_class.Cr)
    KeyFrame.Cb = encoder(Frame_class.Cb)
    return KeyFrame


def decompress_KeyFrame(KeyFrame, decoder, subsampling):
    Y = decoder(KeyFrame.Y)
    Cr = decoder(KeyFrame.Cr)
    Cb = decoder(KeyFrame.Cb)
    ## TO DO
    frame_image = frame_layers_to_image(Y, Cr, Cb, subsampling)
    return frame_image


def compress_not_KeyFrame(Frame_class, KeyFrame, encoder, compress_rate):
    Compress_data = data()
    ## TO DO
    Compress_data.Y =(Frame_class.Y - KeyFrame.Y) // compress_rate
    Compress_data.Cb = (Frame_class.Cb - KeyFrame.Cb) // compress_rate
    Compress_data.Cr = (Frame_class.Cr - KeyFrame.Cr) // compress_rate
    Compress_data.Y = encoder(Compress_data.Y)
    Compress_data.Cr = encoder(Compress_data.Cr)
    Compress_data.Cb = encoder(Compress_data.Cb)
    return Compress_data


def decompress_not_KeyFrame(Compress_data, KeyFrame, decoder, compress_rate, subsampling):
    ## TO DO
    Y = decoder(Compress_data.Y)
    Y = (Y * compress_rate) + KeyFrame.Y
    Cr = decoder(Compress_data.Cr)
    Cr = (Cr * compress_rate) + KeyFrame.Cr
    Cb = decoder(Compress_data.Cb)
    Cb = (Cb * compress_rate) + KeyFrame.Cb

    return frame_layers_to_image(Y, Cr, Cb, subsampling)


def plotDiffrence(ReferenceFrame, DecompressedFrame, ROI, subsampling, dzielnik, method="None"):
    # bardzo słaby i sztuczny przykład wykorzystania tej opcji
    # przerobić żeby porównanie było dokonywane w RGB nie YCrCb i/lub zastąpić innym porównaniem
    # ROI - Region of Insert współrzędne fragmentu który chcemy przybliżyć i ocenić w formacie [w1,w2,k1,k2]

    # fig, axs = plt.subplots(1, 3, sharey=True)
    # fig.set_size_inches(16, 5)
    #
    # axs[0].imshow(ReferenceFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]])
    # axs[2].imshow(DecompressedFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]])
    # diff = ReferenceFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float) - DecompressedFrame[ROI[0]:ROI[1],
    #                                                                     ROI[2]:ROI[3]].astype(float)
    # print(np.min(diff), np.max(diff))
    # axs[1].imshow(diff, vmin=np.min(diff), vmax=np.max(diff))
    # plt.show()

    ReferenceFrame_RGB = cv2.cvtColor(ReferenceFrame, cv2.COLOR_YCrCb2RGB)
    DecompressedFrame_RGB = cv2.cvtColor(DecompressedFrame, cv2.COLOR_YCrCb2RGB)

    ref_roi = ReferenceFrame_RGB[ROI[0]:ROI[1], ROI[2]:ROI[3]]
    decomp_roi = DecompressedFrame_RGB[ROI[0]:ROI[1], ROI[2]:ROI[3]]

    diff = ref_roi.astype(float) - decomp_roi.astype(float)
    print(np.min(diff), np.max(diff))

    diff_min = diff.min()
    diff_max = diff.max()

    if diff_max != diff_min:
        diff_normalized = (diff - diff_min) / (diff_max - diff_min) * 255
        diff_normalized = diff_normalized.astype(np.uint8)
    else:
        diff_normalized = np.zeros_like(diff, dtype=np.uint8)

    fig, axs = plt.subplots(1, 3, sharey=True)
    fig.set_size_inches(16, 5)

    axs[0].imshow(ref_roi)
    axs[0].set_title("Reference Frame (RGB)")
    axs[1].imshow(diff_normalized)
    axs[1].set_title("Difference (Normalized)")
    axs[2].imshow(decomp_roi)
    axs[2].set_title("Decompressed Frame (RGB)")
    if method=="None":
        plt.suptitle("File:{}, subsampling={}, divider={}, KeyFrame={} ".format(plik, subsampling, dzielnik, key_frame_counter))
    else:
        plt.suptitle("File:{}, subsampling={}, divider={}, KeyFrame={}, Method=ByteRun".format(plik, subsampling, dzielnik, key_frame_counter))
    # plt.show()

    # Save plot as an image
    safe_subsampling = subsampling.replace(":", "_")
    plot_filename = os.path.join(temp_dir, f"plot_{safe_subsampling}_{dzielnik}_{method}.png")
    plt.savefig(plot_filename)
    plt.close(fig)

    # Add plot to Word document
    doc.add_heading(f"Subsampling: {subsampling}, Divider: {dzielnik}, Method: {method}", level=2)
    doc.add_picture(plot_filename, width=Inches(6))
##############################################################################
####     Głowna pętla programu      ##########################################
##############################################################################

doc = Document()
doc.add_heading("Compression Results", level=1)

# Directory for temporary plot images
temp_dir = "temp_plots"
os.makedirs(temp_dir, exist_ok=True)

# Zadanie 1
for plik in pliki:
    for subsampling in subsampling_rates:
        for dzielnik in dzielniki:
            if plik == "clip_5.mp4":
                ROI = ROI5
            cap = cv2.VideoCapture(os.path.join(kat, plik))
            if ile < 0:
                ile = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            # cv2.namedWindow('Normal Frame')
            # cv2.namedWindow('Decompressed Frame')
            compression_information = np.zeros((3, ile))
            for i in range(ile):
                ret, frame = cap.read()
                if wyswietlaj_kaltki:
                    cv2.imshow('Normal Frame', frame)
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
                Frame_class = frame_image_to_class(frame, subsampling)
                if (i % key_frame_counter) == 0:  # pobieranie klatek kluczowych
                    KeyFrame = compress_KeyFrame(Frame_class, no_encoder)
                    cY = KeyFrame.Y
                    cCb = KeyFrame.Cb
                    cCr = KeyFrame.Cr
                    Decompresed_Frame = decompress_KeyFrame(KeyFrame, no_decoder, subsampling)
                else:  # kompresja,
                    Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, no_encoder, dzielnik)
                    cY = Compress_data.Y
                    cCb = Compress_data.Cb
                    cCr = Compress_data.Cr
                    Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame, no_decoder, dzielnik, subsampling)

                compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
                compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
                compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size
                if wyswietlaj_kaltki:
                    cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

                if np.any(plot_frames == i):  # rysuj wykresy
                    for r in ROI:
                        plotDiffrence(frame, Decompresed_Frame, r, subsampling, dzielnik)

                if np.any(auto_pause_frames == i):
                    cv2.waitKey(-1)  # wait until any key is pressed

                k = cv2.waitKey(1) & 0xff

                if k == ord('q'):
                    break
                elif k == ord('p'):
                    cv2.waitKey(-1)  # wait until any key is pressed

# Zadanie 2
plik = "clip_3.mp4"
for key_frame_counter in key_frame_counters:
    cap = cv2.VideoCapture(os.path.join(kat, plik))
    if ile < 0:
        ile = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    # cv2.namedWindow('Normal Frame')
    # cv2.namedWindow('Decompressed Frame')
    compression_information = np.zeros((3, ile))
    for i in range(ile):
        ret, frame = cap.read()
        if not ret:
            break

        if wyswietlaj_kaltki:
            cv2.imshow('Normal Frame', frame)

        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
        Frame_class = frame_image_to_class(frame, subsampling)

        if (i % key_frame_counter) == 0:
            # Kompresja klatki kluczowej
            KeyFrame = compress_KeyFrame(Frame_class, no_encoder)

            # Kodowanie ByteRun
            cY = ByteRun_encoder(KeyFrame.Y)
            cCb = ByteRun_encoder(KeyFrame.Cb)
            cCr = ByteRun_encoder(KeyFrame.Cr)

            # Dekodowanie ByteRun
            dY = ByteRun_decoder(cY).reshape(KeyFrame.Y.shape)
            dCb = ByteRun_decoder(cCb).reshape(KeyFrame.Cb.shape)
            dCr = ByteRun_decoder(cCr).reshape(KeyFrame.Cr.shape)

            # Zamiana na obiekt data i dekompresja
            DecodedKeyFrame = data()
            DecodedKeyFrame.Y = dY
            DecodedKeyFrame.Cb = dCb
            DecodedKeyFrame.Cr = dCr
            Decompresed_Frame = frame_layers_to_image(DecodedKeyFrame.Y, DecodedKeyFrame.Cr, DecodedKeyFrame.Cb,
                                                      subsampling)

        else:
            # Kompresja nie-kluczowej (różnicowej) klatki
            Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, no_encoder, dzielnik)

            # Kodowanie ByteRun
            cY = ByteRun_encoder(Compress_data.Y)
            cCb = ByteRun_encoder(Compress_data.Cb)
            cCr = ByteRun_encoder(Compress_data.Cr)

            # Dekodowanie ByteRun
            dY = ByteRun_decoder(cY).reshape(Compress_data.Y.shape)
            dCb = ByteRun_decoder(cCb).reshape(Compress_data.Cb.shape)
            dCr = ByteRun_decoder(cCr).reshape(Compress_data.Cr.shape)

            # Zamiana na obiekt data
            DecodedCompressData = data()
            DecodedCompressData.Y = dY
            DecodedCompressData.Cb = dCb
            DecodedCompressData.Cr = dCr

            # Odtworzenie klatki różnicowej
            Decompresed_Frame = decompress_not_KeyFrame(DecodedCompressData, KeyFrame, no_decoder, dzielnik,
                                                        subsampling)

        # Informacje o kompresji
        compression_information[0, i] = (frame[:, :, 0].size - cY.size) / frame[:, :, 0].size
        compression_information[1, i] = (frame[:, :, 0].size - cCb.size) / frame[:, :, 0].size
        compression_information[2, i] = (frame[:, :, 0].size - cCr.size) / frame[:, :, 0].size
        if wyswietlaj_kaltki:
            cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

        if np.any(plot_frames == i):
            for r in ROI:
                plotDiffrence(frame, Decompresed_Frame, r, subsampling, dzielnik, method="ByteRun")

        if np.any(auto_pause_frames == i):
            cv2.waitKey(-1)

        k = cv2.waitKey(1) & 0xff
        if k == ord('q'):
            break
        elif k == ord('p'):
            cv2.waitKey(-1)

# Save the Word document
doc.save("results.docx")

# Clean up temporary files
for file in os.listdir(temp_dir):
    os.remove(os.path.join(temp_dir, file))
os.rmdir(temp_dir)