import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
import scipy.fftpack

# czesc pierwsza
img1 = plt.imread('IMG_INTRO/A1.png')
# print(img1.dtype)
# print(img1.shape)
# print(np.min(img1), np.max(img1))


def imgToUInt8(img):
    if np.issubdtype(img.dtype, np.unsignedinteger):
        return img
    else:
        return (img*255).astype(np.uint8)

def imgToFloat(img):
    if np.issubdtype(img.dtype, np.floating):
        return img
    else:
        return img/255.0



# czesc druga
# matplotlib
# R = img1[:,:,0]
# G = img1[:,:,1]
# B = img1[:,:,2]
# Y1 = 0.299*R + 0.587*G + 0.114*B
# Y2 = 0.2126*R + 0.7152*G + 0.0722*B
# plt.imshow(Y1, cmap=plt.cm.gray)
# plt.show()

# opencv
# img2 = cv2.imread('IMG_INTRO/A1.png')
# img_RGB = cv2.cvtColor(img_BGR, cv2.COLOR_BGR2RGB)
# img_BGR = cv2.cvtColor(img_RGB, cv2.COLOR_RGB2BGR)
# plt.imshow(img2)
# plt.show()


# czesc trzecia
def zad2(img, memfile):
    plt.figure(figsize=(10, 10))  # Adjust the figure size if needed

    plt.subplot(3, 3, 1)
    plt.imshow(img)
    plt.title('Original')

    # Y1 i Y2
    plt.subplot(3, 3, 2)
    Y1 = 0.299 * img[:, :, 0] + 0.587 * img[:, :, 1] + 0.114 * img[:, :, 2]
    plt.imshow(Y1, cmap=plt.cm.gray)
    plt.title('Y1')

    plt.subplot(3, 3, 3)
    Y2 = 0.2126 * img[:, :, 0] + 0.7152 * img[:, :, 1] + 0.0722 * img[:, :, 2]
    plt.imshow(Y2, cmap=plt.cm.gray)
    plt.title('Y2')

    # RGB
    plt.subplot(3, 3, 4)
    plt.imshow(img[:, :, 0], cmap=plt.cm.gray)
    plt.title('Red with cmap=gray')

    plt.subplot(3, 3, 5)
    plt.imshow(img[:, :, 1], cmap=plt.cm.gray)
    plt.title('Green with cmap=gray')

    plt.subplot(3, 3, 6)
    plt.imshow(img[:, :, 2], cmap=plt.cm.gray)
    plt.title('Blue with cmap=gray')

    # Zerowanie kolorow
    plt.subplot(3, 3, 7)
    img2 = img.copy()
    img2[:, :, 1] = 0
    img2[:, :, 2] = 0
    plt.imshow(img2)
    plt.title('Blue = 0, Green = 0')

    plt.subplot(3, 3, 8)
    img2 = img.copy()
    img2[:, :, 0] = 0
    img2[:, :, 2] = 0
    plt.imshow(img2)
    plt.title('Blue = 0, Red = 0')

    plt.subplot(3, 3, 9)
    img2 = img.copy()
    img2[:, :, 0] = 0
    img2[:, :, 1] = 0
    plt.imshow(img2)
    plt.title('Red = 0, Green = 0')

    plt.tight_layout(pad=1.2)
    plt.savefig(memfile, format='png')
    memfile.seek(0)
    plt.show()


# img3 = plt.imread('IMG_INTRO/B01.png')
# zad2(img3)


# czesc czwarta
image_files = ['IMG_INTRO/B01.png', 'IMG_INTRO/B02.jpg']

fragments_coords = [
    (0, 0, 200, 200),
    (450, 450, 650, 650),
    (700, 700, 900, 900)
]


# Zapis w .docx
document = Document()
document.add_heading('lab1', 0)
for image_file in image_files:
    img = plt.imread(image_file)
    for i, (w1, k1, w2, k2) in enumerate(fragments_coords):
        fragment = img[w1:w2, k1:k2].copy()
        memfile = BytesIO()
        zad2(fragment, memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()

document.save('raport.docx')  # zapis do pliku