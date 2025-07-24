import numpy as np
import matplotlib.pyplot as plt
import matplotlib
import cv2
import scipy.fftpack
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
import sys

class Ver2:
    def __init__(self,Y,Cb,Cr,OGShape,Ratio="4:4:4",QY=np.ones((8,8)),QC=np.ones((8,8))):
        self.shape = OGShape
        self.Y=Y
        self.Cb=Cb
        self.Cr=Cr
        self.ChromaRatio=Ratio
        self.QY=QY
        self.QC=QC

def get_size(obj, seen=None):
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    seen.add(obj_id)
    if isinstance(obj, np.ndarray):  # Correct type check for numpy arrays
        size = obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size
def dct2(a):
    return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
    return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B

def RLE_encode(data):
    original_shape = data.shape
    data = data.copy().flatten()
    out_data = np.zeros(len(data) * 2, dtype=int)
    cnt = 1
    index = 0
    for i in range(1, len(data)):
        if data[i] == data[i - 1]:
            cnt += 1
        else:
            out_data[index] = cnt
            out_data[index + 1] = data[i - 1]
            cnt = 1
            index += 2
    out_data[index] = cnt
    out_data[index + 1] = data[-1]
    return np.concatenate(([len(original_shape)], original_shape, out_data[:index + 2]))

def RLE_decode(data):
    shape_length = int(data[0])
    original_shape = tuple(data[1:1 + shape_length].astype(int))
    encoded_data = data[1 + shape_length:]
    original_size = int(np.sum(encoded_data[::2]))
    out_data = np.zeros(original_size)
    index = 0
    for i in range(0, len(encoded_data), 2):
        for j in range(int(encoded_data[i])):
            out_data[index] = encoded_data[i + 1]
            index += 1
    return out_data.reshape(original_shape)

def CompressBlock(block, Q):
    block = block - 128
    dct_block = dct2(block)
    quantized = np.round(dct_block / Q).astype(int)
    vector = zigzag(quantized)
    return vector

def DecompressBlock(vector, Q):
    quantized = zigzag(vector)
    dct_block = quantized*Q
    block = idct2(dct_block)
    block = block+128
    return block

def CompressLayer(L,Q):
    S=np.array([])
    for w in range(0,L.shape[0],8):
        for k in range(0,L.shape[1],8):
            block=L[w:(w+8),k:(k+8)]
            S=np.append(S, CompressBlock(block,Q))
    return S

def DecompressLayer(S,Q):
    L= np.zeros((128,128), dtype=np.uint8)
    for idx,i in enumerate(range(0,S.shape[0],64)):
        vector=S[i:(i+64)]
        m=L.shape[1]/8
        k=int((idx%m)*8)
        w=int((idx//m)*8)
        L[w:(w+8),k:(k+8)]=DecompressBlock(vector,Q)
    return L

def chroma_subsampling(A, Ratio):
    B = np.copy(A)
    if Ratio == "4:2:2":
        B = B[:,::2]
    elif Ratio == "4:2:0":
        pass
    else:
        pass
    return B

def chroma_resampling(A, Ratio):
    B = np.copy(A)
    if Ratio == "4:2:2":
        B = B[B != 0].reshape(128, 64)
        B = np.repeat(B, 2, axis=1).reshape(128, 128)
    return B

# Kompresja
def JPEG_compress(RGB, Ratio="4:4:4", QY=np.ones((8, 8)), QC=np.ones((8, 8))):
    YCbCr = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    Y = YCbCr[:, :, 0]
    Cr = YCbCr[:, :, 1]
    Cb = YCbCr[:, :, 2]
    JPEG = Ver2(Y, Cb, Cr, RGB.shape, Ratio, QY, QC)

    # chroma subsampling
    JPEG.Cr = chroma_subsampling(JPEG.Cr, JPEG.ChromaRatio)
    JPEG.Cb = chroma_subsampling(JPEG.Cb, JPEG.ChromaRatio)

    # kompresja warstwy
    JPEG.Y = CompressLayer(JPEG.Y,JPEG.QY)
    JPEG.Cr = CompressLayer(JPEG.Cr,JPEG.QC)
    JPEG.Cb = CompressLayer(JPEG.Cb,JPEG.QC)

    # kompresja bezstratna
    JPEG.Y = RLE_encode(JPEG.Y)
    JPEG.Cr = RLE_encode(JPEG.Cr)
    JPEG.Cb = RLE_encode(JPEG.Cb)
    return JPEG

# Dekompresja
def JPEG_decompress(JPEG):
    # dekompresja bezstratna
    Y = RLE_decode(JPEG.Y)
    Cr = RLE_decode(JPEG.Cr)
    Cb = RLE_decode(JPEG.Cb)

    # dekompresja warstwy
    Y = DecompressLayer(Y, JPEG.QY)
    Cr = DecompressLayer(Cr, JPEG.QC)
    Cb = DecompressLayer(Cb, JPEG.QC)

    # chroma resampling
    Cr = chroma_resampling(Cr, JPEG.ChromaRatio)
    Cb = chroma_resampling(Cb, JPEG.ChromaRatio)

    # rekonstrukcja obrazu
    YCrCb = np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)

    # YCrCb -> RGB
    RGB = cv2.cvtColor(YCrCb, cv2.COLOR_YCrCb2RGB)
    return RGB

QY= np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])

QC= np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])

QN= np.ones((8,8))

df = pd.DataFrame(data={
    'Zdjecia': ['obraz1.jpg', 'obraz2.jpg', 'obraz3.jpg','obraz4.jpg'],
    "Fragmenty":[
        [[0,200],[1200,400],[1200,100]],
        [[1000,800],[500,400],[1000,500]],
        [[600,480],[500,200],[760,400]],
        [[0,580],[1620,800],[930,540]]
    ]})

# matplotlib.use("TkAgg")
# img = cv2.imread('obraz3.jpg')
# img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
# plt.imshow(img)
# plt.show()

quants = [["quantization table", QY, QC], ["ones table", QN, QN]]
ratios = ["4:4:4", "4:2:2"]
document = Document()

for index, row in df.iterrows():
    # Wgranie calego zdjecia
    document.add_heading(f'Image: {row["Zdjecia"]}', level=1)
    img = cv2.imread(row['Zdjecia'])
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    memfile = BytesIO()
    plt.imshow(img)
    plt.axis('off')
    plt.title('Original image')
    plt.savefig(memfile, format='png', bbox_inches='tight')
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(6))
    memfile.close()
    plt.close()
    height, width, _ = img.shape
    for fragment in row['Fragmenty']:
        x_start, y_start = fragment
        x_end = min(x_start + 128, width)
        y_end = min(y_start + 128, height)
        img_fragment = img[y_start:y_end, x_start:x_end].copy()
        for quant in quants:
            for ratio in ratios:
                fig, axs = plt.subplots(4, 2, sharey=True)
                fig.suptitle(f'Fragment: {fragment}, Ratio: {ratio}, Quant: {quant[0]}')
                fig.set_size_inches(9, 13)

                # obraz oryginalny
                axs[0, 0].imshow(img_fragment)  # RGB
                img_before_YCrCb = cv2.cvtColor(img_fragment, cv2.COLOR_RGB2YCrCb)
                axs[1, 0].imshow(img_before_YCrCb[:, :, 0], cmap=plt.cm.gray)
                axs[1, 0].set_title('Y Channel (Original)')
                axs[2, 0].imshow(img_before_YCrCb[:, :, 1], cmap=plt.cm.gray)
                axs[2, 0].set_title('Cr Channel (Original)')
                axs[3, 0].imshow(img_before_YCrCb[:, :, 2], cmap=plt.cm.gray)
                axs[3, 0].set_title('Cb Channel (Original)')

                # kompresja i dekompresja
                compressed = JPEG_compress(img_fragment, Ratio=ratio, QY=quant[1], QC=quant[2])
                decompressed = JPEG_decompress(compressed)

                # obliczanie stopnia kompresji
                size_Y_before = get_size(img_before_YCrCb[:,:, 0])
                size_Cb_before = get_size(img_before_YCrCb[:, :, 1])
                size_Cr_before = get_size(img_before_YCrCb[:, :, 2])
                size_Y_after = get_size(compressed.Y)
                size_Cb_after = get_size(compressed.Cb)
                size_Cr_after = get_size(compressed.Cr)

                # obraz po dekompresji
                axs[0, 1].imshow(decompressed)  # RGB
                img_after_YCrCb = cv2.cvtColor(decompressed, cv2.COLOR_RGB2YCrCb)
                axs[1, 1].imshow(img_after_YCrCb[:, :, 0], cmap=plt.cm.gray)
                axs[1, 1].set_title('Y Channel (Decompressed) - Compression Ratio: {:.2f}%'.format(size_Y_after / size_Y_before * 100))
                axs[2, 1].imshow(img_after_YCrCb[:, :, 1], cmap=plt.cm.gray)
                axs[2, 1].set_title('Cr Channel (Decompressed) - Compression Ratio: {:.2f}%'.format(size_Cr_after / size_Cr_before * 100))
                axs[3, 1].imshow(img_after_YCrCb[:, :, 2], cmap=plt.cm.gray)
                axs[3, 1].set_title('Cb Channel (Decompressed) - Compression Ratio: {:.2f}%'.format(size_Cb_after / size_Cb_before * 100))

                # Save the plot to a memory buffer
                memfile = BytesIO()
                fig.tight_layout()
                fig.savefig(memfile, format='png')
                memfile.seek(0)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()

                plt.close(fig)
document.save('Sprawozdanie.docx')