import numpy as np
import matplotlib.pyplot as plt
import sys
import cv2
from tqdm import tqdm
from PIL import Image
import os
from docx import Document
from docx.shared import Inches

def get_size(obj, seen=None):
    size = sys.getsizeof(obj)
    if seen is None:
        seen=set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    seen.add(obj_id)
    if isinstance(obj_id):
        size = obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v,seen) for v in obj.values()])
        size += sum([get_size(k,seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size


def count_same_symbols(obj):
    number_of_symbols = 1
    for i in range(len(obj)-1):
        if obj[i] == obj[i+1]:
            number_of_symbols += 1
        else:
            break
    return number_of_symbols

def count_diff_symbols(obj):
    number_of_symbols = 1
    for i in range(len(obj) - 1):
        if obj[i] == obj[i + 1]:
            number_of_symbols -= 1
            break
        else:
            number_of_symbols += 1
    return number_of_symbols

def RLE_encoder(obj):
    obj_temp = obj.copy()
    obj_temp = np.array(obj_temp)
    if obj_temp.shape != (1,):
        obj_temp = obj_temp.flatten()
    encoded = []
    i=0
    while i < len(obj_temp):
        num_of_symbols = count_same_symbols(obj_temp[i:])
        encoded.append(num_of_symbols)
        encoded.append(obj_temp[i])
        i += num_of_symbols
    return encoded

def RLE_decoder(encoded):
    decoded = []
    for i in range(len(encoded)):
        if i % 2 == 0:
            for j in range(encoded[i]):
                decoded.append(encoded[i+1])
    return decoded

def ByteRun_encoder(obj):
    obj_temp = obj.copy()
    obj_temp = np.array(obj_temp)
    if obj_temp.shape != (1,):
        obj_temp = obj_temp.flatten()
    encoded = []
    i = 0
    while i < len(obj_temp):
        num_of_symbols = count_same_symbols(obj_temp[i:])
        if num_of_symbols > 1:
            encoded.append(-num_of_symbols + 1)
            encoded.append(obj_temp[i])
            i += num_of_symbols
        else:
            num_of_symbols = count_diff_symbols(obj_temp[i:])
            encoded.append(num_of_symbols-1)
            encoded.extend(obj_temp[i:i + num_of_symbols])
            i += num_of_symbols
    return encoded

def ByteRun_decoder(encoded):
    decoded = []
    i = 0
    while i < len(encoded):
        if encoded[i] < 0:  # Repeated symbols
            count = -encoded[i] + 1
            symbol = encoded[i + 1]
            decoded.extend([symbol] * count)
            i += 2
        else:  # Unique symbols
            count = encoded[i] + 1
            decoded.extend(encoded[i + 1:i + 1 + count])
            i += 1 + count
    return decoded

def test_cases_compression(data, method_name, encoder, decoder, test_number):
    print(f"Testing {method_name} compression...")
    compressed = encoder(data)
    decompressed = decoder(compressed)
    # doc.add_heading(f"{test_number} - {method_name}", level=2)
    # doc.add_paragraph(f"Original data - {data}")
    # doc.add_paragraph(f"Compressed data - {compressed}")
    # doc.add_paragraph(f"Decompressed data - {decompressed}")

def test_compression(image_data, method_name, encoder, decoder, image_name):
    print(f"Testing {image_name} with {method_name} compression...")

    # Kompresja
    compressed = encoder(image_data)
    # Dekompresja
    decompressed = decoder(compressed)

    # Miary kompresji
    CR = len(image_data) / len(compressed)
    PR = (len(compressed) / len(image_data)) * 100

    # Check if decompression is successful
    success = np.array_equal(image_data, decompressed)
    result_text = f"{method_name} Kompresja {'udana' if success else 'nieudana'}"
    print(result_text)

    # decompressed_image = np.array(decompressed, dtype=image_data.dtype).reshape(image_data.shape)
    # cv2.imwrite(f"C:/Users/smh2k/Desktop/{image_name}_{method_name}_decompressed.png", decompressed_image)

    # # Add results to Word document
    # doc.add_heading(f"{image_name} - {method_name}", level=2)
    # doc.add_paragraph(result_text)
    # doc.add_paragraph(f"Stopień kompresji: {CR:.2f}")
    # doc.add_paragraph(f"Procent kompresji: {PR:.2f}%")

# Example usage
# doc = Document()
# doc.add_heading("Compression Test Results", level=1)

image1 = cv2.imread("schemat.jpg")
image2 = cv2.imread("pkk.jpg")
image3 = cv2.imread("kolorowe zdjecie.jpg")
images = {
    "Rysunek techniczny": image1,
    "Wzór dokumentu": image2,
    "Kolorowe zdjecie": image3
}
test_cases = {
    "Test case 1:" : np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),
    "Test case 2:" : np.array([1,2,3,1,2,3,1,2,3]),
    "Test case 3:" : np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),
    "Test case 4:" : np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]),
    "Test case 5:" : np.zeros((1,520)),
    "Test case 6:" : np.arange(0,521,1),
    "Test case 7:" :np.eye(7),
    "Test case 8:" : np.dstack([np.eye(7),np.eye(7),np.eye(7)]),
    "Test case 9:" : np.ones((1,1,1,1,1,1,10))
}

# for test_number, test_data in test_cases.items():
    # test_cases_compression(test_data, "RLE", RLE_encoder, RLE_decoder, test_number)
    # test_cases_compression(test_data, "ByteRun", ByteRun_encoder, ByteRun_decoder, test_number)

for image_name, image_data in images.items():
    print(f"Processing {image_name}...")
    image_data = image_data.flatten()  # Spłaszczenie obrazu do 1D
    test_compression(image_data, "RLE", RLE_encoder, RLE_decoder, image_name)
    test_compression(image_data, "ByteRun", ByteRun_encoder, ByteRun_decoder, image_name)

# Save the Word document
# doc.save("Sprawozdanie z laboratorium nr6.docx")